# accounts/file_utils.py

import logging
import os

logger = logging.getLogger(__name__)

MAX_FILE_SIZE_MB = 15


class FileExtractionError(Exception):
    """خطای قابل‌نمایش به کاربر هنگام استخراج متن از فایل."""
    pass


def extract_text_from_file(file) -> str:
    if not file:
        raise FileExtractionError("فایلی ارسال نشده است.")

    file.seek(0, os.SEEK_END)
    size = file.tell()
    file.seek(0)

    if size == 0:
        raise FileExtractionError("فایل ارسالی خالی است.")

    if size > MAX_FILE_SIZE_MB * 1024 * 1024:
        raise FileExtractionError(f"حجم فایل بیشتر از {MAX_FILE_SIZE_MB} مگابایت مجاز است.")

    ext = os.path.splitext(file.name)[1].lower().lstrip(".")

    try:
        if ext == "txt":
            text = _extract_txt(file)
        elif ext == "pdf":
            text = _extract_pdf(file)
        elif ext == "docx":
            text = _extract_docx(file)
        elif ext == "doc":
            text = _extract_doc(file)
        elif ext == "rtf":
            text = _extract_rtf(file)
        else:
            raise FileExtractionError(
                "فرمت فایل پشتیبانی نمی‌شود. لطفاً فایل txt، pdf، docx یا doc آپلود کنید."
            )
    except FileExtractionError:
        raise
    except Exception as e:
        logger.exception("خطای غیرمنتظره در استخراج متن از فایل: %s", file.name)
        raise FileExtractionError(
            "فایل ارسالی قابل باز شدن نیست یا خراب است. لطفاً فایل دیگری امتحان کنید."
        ) from e

    text = (text or "").strip()

    if not text:
        raise FileExtractionError(
            "متنی از فایل استخراج نشد. اگر فایل اسکن‌شده یا تصویری است، "
            "لطفاً نسخه‌ی متنی (Word یا PDF قابل کپی) آپلود کنید."
        )

    return text


def _extract_txt(file) -> str:
    raw = file.read()
    encodings = ["utf-8", "utf-8-sig", "cp1256", "windows-1252"]
    for enc in encodings:
        try:
            return raw.decode(enc)
        except (UnicodeDecodeError, LookupError):
            continue
    return raw.decode("utf-8", errors="ignore")


def _reshape_bidi(text: str) -> str:
    """
    تبدیل متن فارسی/عربی از حالت منطقی (logical order) که pdfplumber
    استخراج می‌کند، به ترتیب صحیح نمایشی با استفاده از الگوریتم استاندارد
    Unicode Bidirectional Algorithm.
    """
    try:
        import arabic_reshaper
        from bidi.algorithm import get_display
    except ImportError:
        raise FileExtractionError(
            "پردازش متن فارسی/عربی PDF نیاز به کتابخانه دارد. "
            "لطفاً با پشتیبانی تماس بگیرید (نیاز به نصب: "
            "pip install python-bidi arabic-reshaper)."
        )

    fixed_lines = []
    for line in text.split("\n"):
        if not line.strip():
            fixed_lines.append(line)
            continue
        try:
            reshaped = arabic_reshaper.reshape(line)
            display_line = get_display(reshaped)
            fixed_lines.append(display_line)
        except Exception:
            fixed_lines.append(line)

    return "\n".join(fixed_lines)


def _extract_pdf(file) -> str:
    try:
        import pdfplumber
    except ImportError:
        raise FileExtractionError(
            "پردازش فایل PDF در سرور فعال نیست. لطفاً با پشتیبانی تماس بگیرید "
            "(نیاز به نصب pdfplumber: pip install pdfplumber)."
        )

    text_parts = []
    try:
        with pdfplumber.open(file) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text() or ""
                text_parts.append(page_text)
    except Exception as e:
        logger.exception("خطا در باز کردن فایل PDF: %s", file.name)
        raise FileExtractionError(
            "فایل PDF قابل باز شدن نیست، خراب است یا رمزگذاری‌شده است."
        ) from e

    text = "\n".join(text_parts).strip()

    if not text:
        text = _try_ocr_pdf(file)
        return text

    text = _reshape_bidi(text)

    return text


def _try_ocr_pdf(file) -> str:
    """
    OCR برای فایل‌های PDF اسکن‌شده/تصویری.
    نیازمند نصب روی سرور:
        pip install pdf2image pytesseract
        apt-get install poppler-utils tesseract-ocr tesseract-ocr-fas
    """
    try:
        import pytesseract
        from pdf2image import convert_from_bytes
    except ImportError:
        return ""

    try:
        file.seek(0)
        images = convert_from_bytes(file.read())
        text_parts = []
        for img in images:
            page_text = pytesseract.image_to_string(img, lang="fas+eng")
            text_parts.append(page_text)
        return "\n".join(text_parts).strip()
    except Exception:
        logger.exception("خطا در OCR فایل PDF")
        return ""


def _extract_docx(file) -> str:
    try:
        import docx
    except ImportError:
        raise FileExtractionError(
            "پردازش فایل Word (docx) در سرور فعال نیست. لطفاً با پشتیبانی تماس بگیرید "
            "(نیاز به نصب python-docx: pip install python-docx)."
        )

    try:
        document = docx.Document(file)
    except Exception as e:
        logger.exception("خطا در باز کردن فایل docx: %s", file.name)
        raise FileExtractionError(
            "فایل Word قابل باز شدن نیست یا خراب است."
        ) from e

    parts = []

    for para in document.paragraphs:
        if para.text.strip():
            parts.append(para.text)

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    parts.append(cell.text)

    return "\n".join(parts)


def _extract_doc(file) -> str:
    try:
        import textract
    except ImportError:
        raise FileExtractionError(
            "پردازش فایل Word قدیمی (doc) در سرور فعال نیست. لطفاً با پشتیبانی تماس بگیرید "
            "(نیاز به نصب textract: pip install textract)."
        )

    try:
        file.seek(0)
        raw_bytes = file.read()
        tmp_path = "/tmp/_upload_tmp.doc"
        with open(tmp_path, "wb") as f:
            f.write(raw_bytes)
        text_bytes = textract.process(tmp_path)
        return text_bytes.decode("utf-8", errors="ignore")
    except Exception as e:
        logger.exception("خطا در پردازش فایل doc: %s", file.name)
        raise FileExtractionError(
            "فایل Word (doc) قابل باز شدن نیست یا خراب است."
        ) from e
    finally:
        try:
            os.remove(tmp_path)
        except Exception:
            pass


def _extract_rtf(file) -> str:
    try:
        from striprtf.striprtf import rtf_to_text
    except ImportError:
        raise FileExtractionError(
            "پردازش فایل RTF در سرور فعال نیست. لطفاً با پشتیبانی تماس بگیرید "
            "(نیاز به نصب striprtf: pip install striprtf)."
        )

    try:
        file.seek(0)
        raw = file.read()
        try:
            raw_text = raw.decode("utf-8")
        except UnicodeDecodeError:
            raw_text = raw.decode("cp1256", errors="ignore")
        return rtf_to_text(raw_text)
    except Exception as e:
        logger.exception("خطا در پردازش فایل rtf: %s", file.name)
        raise FileExtractionError(
            "فایل RTF قابل باز شدن نیست یا خراب است."
        ) from e
